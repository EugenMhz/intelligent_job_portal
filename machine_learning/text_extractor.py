"""
text_extractor.py
------------------
Step 1 & 2 of the pipeline (see Chapter 4 pseudocode: handleResumeUpload):
    - saveResumeToDatabase()  -> handled elsewhere (backend/DB team)
    - ResumeParser.extractText(resumeFile)  -> THIS FILE

Purpose:
    Given a resume file (PDF or DOCX), pull out the raw, unstructured text
    so it can be cleaned and analyzed later.

Why two extractors?
    Resumes come in different formats. PDFs and Word docs store text very
    differently internally, so each needs its own extraction method.
"""

import os
import pdfplumber
import docx  # from python-docx


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extracts all text from a PDF file, page by page.

    pdfplumber reads the PDF's internal text layer directly (it does NOT
    do OCR / image recognition). This works well for resumes that were
    exported from Word/Canva/LinkedIn etc., but will return little or
    nothing for scanned image-only PDFs.
    """
    full_text = []

    with pdfplumber.open(file_path) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text()
            if page_text:  # page_text can be None if page has no text layer
                full_text.append(page_text)
            else:
                print(f"[warning] No extractable text found on page {page_number}. "
                      f"This page may be an image/scan.")

    return "\n".join(full_text)


def extract_text_from_docx(file_path: str) -> str:
    """
    Extracts text from a .docx file by reading each paragraph.

    Note: this does NOT read text inside tables. Many resumes use tables
    for layout (e.g. skills in a 2-column table), so we also pull table
    cell text separately and append it.
    """
    document = docx.Document(file_path)

    full_text = [para.text for para in document.paragraphs if para.text.strip()]

    # Also capture text sitting inside tables (common in resume templates)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)

    return "\n".join(full_text)


def extract_text(file_path: str) -> str:
    """
    Main entry point. Detects the file type from its extension and
    calls the right extractor.

    This is the function the rest of your pipeline (and eventually the
    Flask backend) should call — it doesn't need to know or care whether
    the uploaded file was a PDF or DOCX.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"No such file: {file_path}")

    extension = os.path.splitext(file_path)[1].lower()

    if extension == ".pdf":
        return extract_text_from_pdf(file_path)
    elif extension == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(
            f"Unsupported file type '{extension}'. Only .pdf and .docx are supported."
        )


if __name__ == "__main__":
    # Quick manual test — replace with a real resume path to try it out
    import sys
    if len(sys.argv) > 1:
        text = extract_text(sys.argv[1])
        print(text)
    else:
        print("Usage: python text_extractor.py <path_to_resume.pdf_or_docx>")
